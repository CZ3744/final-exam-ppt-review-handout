from pathlib import Path

from docx import Document

from ppt_review_handout.cli import handout_to_docx


def sample_handout():
    return {
        "chapter_title": "测试章节",
        "source_file": "测试课件.pptx",
        "review_goals": ["掌握核心概念。"],
        "knowledge_framework": ["概念", "对比"],
        "core_points": {"核心知识点": ["这是一个用于测试的知识点。"]},
        "terms": {"测试术语": "用于测试中文字体写入。"},
        "comparison_tables": [
            {
                "title": "宽表格测试",
                "headers": ["项目", "定义", "特点", "用途", "易错点"],
                "rows": [
                    ["材料A", "定义A", "特点A", "用途A", "易错点A"],
                    ["材料B", "定义B", "特点B", "用途B", "易错点B"],
                ],
            }
        ],
        "processes": {"流程": ["步骤1", "步骤2"]},
        "exam_points": ["简答题测试。"],
        "confusing_points": ["易混淆点测试。"],
        "quick_summary": ["速记总结测试。"],
        "slide_count": 2,
        "image_heavy_slides": [],
    }


def test_review_margin_docx_renders_with_wide_table(tmp_path: Path):
    out = tmp_path / "review_margin.docx"
    handout_to_docx(sample_handout(), out, layout="review-margin")
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    # Wide tables in review-margin layout should be converted to vertical cards,
    # so all cell values remain visible in paragraphs instead of being clipped.
    assert "材料A" in text
    assert "定义A" in text
    assert "易错点A" in text


def test_standard_docx_renders_normal_table(tmp_path: Path):
    out = tmp_path / "standard.docx"
    handout_to_docx(sample_handout(), out, layout="standard")
    assert out.exists()
    doc = Document(str(out))
    assert any("宽表格测试" in p.text for p in doc.paragraphs)
    assert len(doc.tables) >= 1
